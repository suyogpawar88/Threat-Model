#!/usr/bin/env python3
"""
Word (.docx) threat model report builder for appsec-threat-modeler.

Takes a structured JSON payload (produced during the threat-modeling skill
workflow -- see skills/threat-modeling/references/report-data-schema.json
for the full shape) and renders a polished Word report: doc info, DFD/threat
model diagram references, STRIDE threat register (with OWASP Top 10 / API
Security Top 10 / LLM Top 10 mapping), risk prioritization, compensating
controls, attack-chain mapping (with MITRE ATT&CK / ATLAS technique tagging),
compliance gap analysis with recommended additional controls, a framework
coverage appendix, and overall recommendations.

Usage:
  python3 build_docx_report.py report_data.json Output_Report.docx
"""
import json
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RISK_COLORS = {
    "High": RGBColor(0xC0, 0x39, 0x2B),
    "Medium": RGBColor(0xB9, 0x77, 0x0E),
    "Low": RGBColor(0x1E, 0x8F, 0x4E),
}
RISK_FILL = {"High": "F8CECC", "Medium": "FFE6CC", "Low": "D5E8D4"}


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_table(doc, headers, rows, widths_cm=None, risk_col_idx=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = "" if val is None else str(val)
            if risk_col_idx is not None and i == risk_col_idx:
                bucket = str(val).split()[0] if val else ""
                for label, fill in RISK_FILL.items():
                    if label.lower() in str(val).lower():
                        _shade_cell(cells[i], fill)
                        break
    if widths_cm:
        _set_col_widths(table, widths_cm)
    doc.add_paragraph()
    return table


def build_report(data: dict, output_path: str):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # --- Title page ---
    title = doc.add_heading(f"Threat Model Report: {data.get('service_name', 'Unnamed Service')}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(
        doc,
        ["Field", "Detail"],
        [
            ["Service", data.get("service_name", "")],
            ["Methodology", data.get("methodology", "STRIDE + attack-chain mapping + compensating-controls assessment + OWASP/MITRE mapping")],
            ["Date", data.get("date", "")],
            ["Status", data.get("status", "Draft")],
            ["Sources", ", ".join(data.get("sources_summary", []))],
            [
                "Risk Summary",
                f"High: {sum(1 for t in data.get('threats', []) if t.get('risk_bucket') == 'High')} | "
                f"Medium: {sum(1 for t in data.get('threats', []) if t.get('risk_bucket') == 'Medium')} | "
                f"Low: {sum(1 for t in data.get('threats', []) if t.get('risk_bucket') == 'Low')}",
            ],
        ],
        widths_cm=[4, 12],
    )
    doc.add_page_break()

    # --- 1. Executive Summary ---
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(data.get("executive_summary", ""))

    # --- 2. Scope and Sources ---
    add_heading(doc, "2. Scope and Sources", 1)
    doc.add_paragraph(data.get("scope", ""))
    sources = data.get("sources", {})
    if sources:
        rows = []
        for system, items in sources.items():
            for item in items:
                rows.append([system.title(), item.get("ref", ""), item.get("description", "")])
        if rows:
            add_table(doc, ["System", "Reference", "Description"], rows, widths_cm=[3, 4, 9])

    # --- 3. Data Flow Diagram ---
    add_heading(doc, "3. Data Flow Diagram (DFD)", 1)
    diagrams = data.get("diagrams", {})
    p = doc.add_paragraph(
        "The Data Flow Diagram and Threat Model Diagram are delivered as separate, editable "
        "draw.io (.drawio) files alongside this report. Open at https://app.diagrams.net or in "
        "the draw.io desktop app."
    )
    if diagrams.get("dfd_image_path"):
        try:
            doc.add_picture(diagrams["dfd_image_path"], width=Cm(16))
        except Exception:
            pass
    doc.add_paragraph(f"DFD file: {diagrams.get('dfd_drawio_filename', 'N/A')}")
    doc.add_paragraph(f"Threat model diagram file (with trust boundaries and threat actors): {diagrams.get('threat_model_drawio_filename', 'N/A')}")

    # --- 4. Actors and Components ---
    add_heading(doc, "4. Actors and Components", 1)
    add_table(
        doc,
        ["Component", "Type", "Description"],
        [[c.get("name"), c.get("type"), c.get("description")] for c in data.get("actors_components", [])],
        widths_cm=[4, 3, 9],
    )

    # --- 5. Trust Boundaries ---
    add_heading(doc, "5. Trust Boundaries", 1)
    add_table(
        doc,
        ["Boundary", "Flows Crossing", "Description"],
        [[b.get("boundary"), b.get("flows_crossing"), b.get("description")] for b in data.get("trust_boundaries", [])],
        widths_cm=[5, 4, 7],
    )

    # --- 6. Data Flows ---
    add_heading(doc, "6. Data Flows", 1)
    add_table(
        doc,
        ["#", "Source", "Destination", "Data / Action", "Protocol", "Threats"],
        [
            [f.get("num"), f.get("source"), f.get("destination"), f.get("data_action"), f.get("protocol"), ", ".join(f.get("threats", []))]
            for f in data.get("data_flows", [])
        ],
        widths_cm=[1, 3, 3, 5, 3, 2],
    )

    # --- 7. STRIDE Threat Register ---
    add_heading(doc, "7. STRIDE Threat Register", 1)
    categories = ["Spoofing", "Tampering", "Repudiation", "Information Disclosure", "Denial of Service", "Elevation of Privilege"]
    threats = data.get("threats", [])
    for idx, cat in enumerate(categories, start=1):
        cat_threats = [t for t in threats if t.get("category") == cat]
        if not cat_threats:
            continue
        add_heading(doc, f"7.{idx} {cat}", 2)
        add_table(
            doc,
            ["ID", "Threat", "Affected Flow", "Threat Actor", "L×I=Score", "Risk", "OWASP", "Mitigation"],
            [
                [
                    t.get("id"),
                    f"{t.get('title', '')} — {t.get('description', '')}",
                    t.get("affected_flow"),
                    t.get("threat_actor"),
                    f"{t.get('likelihood', '')}×{t.get('impact', '')}={t.get('score', '')}",
                    t.get("risk_bucket"),
                    ", ".join(t.get("owasp_id", [])) if isinstance(t.get("owasp_id"), list) else t.get("owasp_id", ""),
                    t.get("mitigation"),
                ]
                for t in cat_threats
            ],
            widths_cm=[1.5, 4.5, 2, 2.5, 1.8, 1.3, 2, 3.4],
            risk_col_idx=5,
        )

    # --- 8. Risk Prioritization Matrix ---
    add_heading(doc, "8. Risk Prioritization Matrix", 1)
    doc.add_paragraph("All threats, sorted by score descending, then likelihood descending.")
    sorted_threats = sorted(threats, key=lambda t: (t.get("score", 0), t.get("likelihood", 0)), reverse=True)
    add_table(
        doc,
        ["Rank", "ID", "Threat", "Score", "Rationale", "Owner Action"],
        [
            [i + 1, t.get("id"), t.get("title"), t.get("score"), t.get("impact_rationale", ""), t.get("owner_action", "")]
            for i, t in enumerate(sorted_threats)
        ],
        widths_cm=[1.5, 1.5, 4, 2, 6, 3],
    )

    # --- 9. Compensating Controls & Mitigation Assurance ---
    add_heading(doc, "9. Compensating Controls & Mitigation Assurance", 1)
    add_table(
        doc,
        ["ID", "Existing Compensating Controls", "Prob. of Successful Mitigation", "Residual Risk", "Additional Controls Required"],
        [
            [
                t.get("id"),
                t.get("existing_controls", "None identified"),
                t.get("probability_of_mitigation", ""),
                t.get("residual_risk", ""),
                "; ".join(t.get("additional_controls_required", [])) if isinstance(t.get("additional_controls_required"), list) else t.get("additional_controls_required", ""),
            ]
            for t in threats
        ],
        widths_cm=[1.5, 5, 3, 2.5, 5],
    )

    # --- 10. Attack Scenario / Attack Chain Mapping ---
    add_heading(doc, "10. Attack Scenario and Attack Chain Mapping", 1)
    doc.add_paragraph(
        "Each attack chain links individual threats from the register into a plausible end-to-end "
        "attack path, showing how a low-severity finding can be combined with others to reach a "
        "high-impact outcome."
    )
    for chain in data.get("attack_chains", []):
        add_heading(doc, f"Attack Chain: {chain.get('name')} ({chain.get('chain_id')})", 2)
        doc.add_paragraph(chain.get("narrative", ""))
        def _mitre_cell(s):
            ids = list(s.get("mitre_attack_technique", []) or []) + list(s.get("mitre_atlas_technique", []) or [])
            return ", ".join(ids)

        add_table(
            doc,
            ["Step", "Threat ID", "Description", "MITRE ATT&CK / ATLAS"],
            [[s.get("step_num"), s.get("threat_id"), s.get("description"), _mitre_cell(s)] for s in chain.get("steps", [])],
            widths_cm=[1.5, 2, 9.5, 3],
        )
        doc.add_paragraph(f"Overall chain risk: {chain.get('overall_risk', '')}")

    # --- 11. Compliance Gap Analysis & Additional Controls ---
    add_heading(doc, "11. Compliance Gap Analysis and Additional Controls", 1)
    doc.add_paragraph(
        "Controls observed to be partially compliant during this assessment, with additional "
        "controls recommended to close the gap to full compliance."
    )
    add_table(
        doc,
        ["Control", "Compliance Status", "Gap Description", "Recommended Additional Controls"],
        [
            [
                g.get("control_name"),
                g.get("compliance_status"),
                g.get("gap_description"),
                "; ".join(g.get("recommended_additional_controls", [])) if isinstance(g.get("recommended_additional_controls"), list) else g.get("recommended_additional_controls", ""),
            ]
            for g in data.get("compliance_gaps", [])
        ],
        widths_cm=[3, 2.5, 6, 6],
    )

    # --- 12. Framework Coverage Mapping ---
    add_heading(doc, "12. Framework Coverage Mapping", 1)
    doc.add_paragraph(
        "Every threat mapped to the applicable OWASP list(s) -- OWASP Top 10 (web AppSec), OWASP "
        "API Security Top 10, and/or OWASP Top 10 for LLM Applications -- and, for AI/ML-targeting "
        "threats, the AI/ML pipeline component involved. See references/owasp-mappings.md and "
        "references/ai-threat-taxonomy.md for the full mapping tables."
    )
    add_table(
        doc,
        ["ID", "Category", "AI Component", "OWASP Mapping"],
        [
            [
                t.get("id"),
                t.get("category"),
                t.get("ai_component", "none"),
                ", ".join(t.get("owasp_id", [])) if isinstance(t.get("owasp_id"), list) else t.get("owasp_id", ""),
            ]
            for t in threats
        ],
        widths_cm=[1.5, 4, 4, 6.5],
    )
    fc = data.get("framework_coverage", {})
    if fc:
        doc.add_paragraph(
            "Frameworks applied to this assessment: "
            + ", ".join(k.replace("_", " ").upper() for k, v in fc.items() if v)
        )

    # --- 13. Key Recommendations ---
    add_heading(doc, "13. Key Recommendations", 1)
    for i, rec in enumerate(data.get("recommendations", []), start=1):
        doc.add_paragraph(f"{i}. {rec}", style="List Number")

    # --- 14. References ---
    add_heading(doc, "14. References", 1)
    for ref in data.get("references", []):
        doc.add_paragraph(ref, style="List Bullet")

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 build_docx_report.py <report_data.json> <output.docx>", file=sys.stderr)
        sys.exit(1)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        data = json.load(f)
    path = build_report(data, sys.argv[2])
    print(f"Wrote {path}")


if __name__ == "__main__":
    main()
